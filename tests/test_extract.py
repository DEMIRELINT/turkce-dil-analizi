from pathlib import Path

import pytest

from dilanaliz.extract import _repair_broken_words, extract_docx, extract_docx_with_report

docx = pytest.importorskip("docx")


class _FakeSpeller:
    """`_repair_broken_words`'ün MANTIĞINI gerçek sözlük olmadan test etmek için
    sahte bir `HunspellChecker` (yalnız `is_known` gerekir)."""

    def __init__(self, known: set[str]) -> None:
        self._known = known

    def is_known(self, word: str) -> bool:
        return word in self._known


def test_repair_merges_when_one_side_invalid_and_union_valid():
    # "değer" geçerli, "lendirme" geçersiz, birleşim "değerlendirme" geçerli.
    speller = _FakeSpeller({"değer", "değerlendirme"})
    text = "Ekip performansını değer lendirme toplantısında tartıştı."
    repaired, n = _repair_broken_words(text, speller)
    assert n == 1
    assert "değerlendirme" in repaired
    assert "değer lendirme" not in repaired


def test_repair_merges_dot_separator():
    speller = _FakeSpeller({"I", "KAPSAMADIĞI"})
    text = "Cihaz bu modu KAPSAMADIĞ.I için manuel ayar gerekir."
    repaired, n = _repair_broken_words(text, speller)
    assert n == 1
    assert "KAPSAMADIĞI" in repaired
    assert "KAPSAMADIĞ.I" not in repaired


def test_repair_skips_when_both_sides_valid():
    # İki taraf da geçerli kelime — deterministik katman dokunmamalı (LLM'e bırakılır).
    speller = _FakeSpeller({"kapı", "dayım", "kapıdayım"})
    text = "Kurye şu anda kapı dayım dedi."
    repaired, n = _repair_broken_words(text, speller)
    assert n == 0
    assert repaired == text


def test_repair_skips_when_union_also_invalid():
    # Birleşim de sözlükte yoksa (asıl kelimede ayrıca karakter bozulması var) dokunma.
    speller = _FakeSpeller({"Değ"})  # "ıştirmeyı" ve "Değıştirmeyı" bilinmiyor
    text = "Sinyal Değ ıştirmeyı düğmesi."
    repaired, n = _repair_broken_words(text, speller)
    assert n == 0
    assert repaired == text


def test_repair_ignores_normal_sentence_boundaries():
    # Nokta + boşluk (normal cümle sonu) tek karakterlik ayraç DEĞİLDİR; dokunulmaz.
    speller = _FakeSpeller({"doğru", "şimdi"})
    text = "Bu doğru. Şimdi devam edelim."
    repaired, n = _repair_broken_words(text, speller)
    assert n == 0
    assert repaired == text


def test_repair_merges_across_newline_gap():
    # Doğrulanmış gerçek kök neden: docx2python, kelime ortasına kaçırılan bir
    # satır içi kırılmayı (<w:br/>) tek "\n" karakteri olarak çıkarır.
    speller = _FakeSpeller({"Değ", "değiştirmeyi", "Değiştirmeyi"})
    text = "Değ\niştirmeyi düğmesi."
    repaired, n = _repair_broken_words(text, speller)
    assert n == 1
    assert "Değiştirmeyi" in repaired
    assert "\n" not in repaired


# --- Gerçek tr_TR sözlüğüyle uçtan uca doğrulama (sözlük yoksa atlanır) -----

_DICT_BASE = "dicts/tr_TR"
_has_dict = Path(f"{_DICT_BASE}.dic").exists()


@pytest.fixture(scope="module")
def real_speller():
    from dilanaliz.spell import HunspellChecker

    return HunspellChecker(_DICT_BASE)


@pytest.mark.skipif(not _has_dict, reason="tr_TR sözlüğü yok")
def test_extract_repairs_split_word_with_real_dict(tmp_path, real_speller):
    path = _make_docx(tmp_path, ["Ekip performansını değer lendirme toplantısında tartıştı."])
    text = extract_docx(path, speller=real_speller)
    assert "değerlendirme" in text
    assert "değer lendirme" not in text


@pytest.mark.skipif(not _has_dict, reason="tr_TR sözlüğü yok")
def test_extract_reports_repaired_word_count(tmp_path, real_speller):
    path = _make_docx(tmp_path, ["Cihaz bu modu KAPSAMADIĞ.I için manuel ayar gerekir."])
    text, report = extract_docx_with_report(path, speller=real_speller)
    assert "KAPSAMADIĞI" in text
    assert report.repaired_words == 1
    assert any("onarıldı" in w for w in report.warnings)


@pytest.mark.skipif(not _has_dict, reason="tr_TR sözlüğü yok")
def test_extract_does_not_merge_two_valid_words(tmp_path, real_speller):
    # "kapı" ve "dayım" ikisi de tek başına geçerli — deterministik katman dokunmamalı.
    path = _make_docx(tmp_path, ["Kurye geldiğinde şu anda kapı dayım dedi."])
    text = extract_docx(path, speller=real_speller)
    assert "kapı dayım" in text


@pytest.mark.skipif(not _has_dict, reason="tr_TR sözlüğü yok")
def test_extract_no_repair_without_speller(tmp_path):
    # speller=None (varsayılan) → davranış tamamen geriye dönük uyumlu.
    path = _make_docx(tmp_path, ["Ekip performansını değer lendirme toplantısında tartıştı."])
    text, report = extract_docx_with_report(path)
    assert "değer lendirme" in text
    assert report.repaired_words == 0


def _add_manual_line_break_mid_word(document, left: str, right: str) -> None:
    """Bir paragrafta LEFT+RIGHT kelimesinin ortasına gerçek bir Word satır içi
    kırılması (`<w:br/>`, Shift+Enter) enjekte eder — docx2python'un bunu nasıl
    çıkardığını uçtan uca doğrulamak için (bkz. `_repair_broken_words` docstring'i:
    bu mekanizma gerçek bir docx ile test edilip doğrulanmıştır)."""
    from docx.oxml import OxmlElement

    paragraph = document.add_paragraph()
    left_run = paragraph.add_run(left)
    left_run._r.append(OxmlElement("w:br"))
    paragraph.add_run(right)


@pytest.mark.skipif(not _has_dict, reason="tr_TR sözlüğü yok")
def test_extract_repairs_real_manual_line_break_mid_word(tmp_path, real_speller):
    # Gerçek docx2python davranışıyla uçtan uca: biçimlendirme farkı (kalın,
    # punto, renk, dil etiketi) TEK BAŞINA sorun YARATMAZ (docx2python düzgün
    # birleştirir) — yalnız gerçek bir <w:br/> satır içi kırılması sorun yaratır.
    document = docx.Document()
    _add_manual_line_break_mid_word(document, "Değ", "iştirmeyi")
    path = tmp_path / "satir_kirilmasi.docx"
    document.save(str(path))

    text, report = extract_docx_with_report(path, speller=real_speller)
    assert "Değiştirmeyi" in text
    assert "\n" not in text.split("\n\n")[0]
    assert report.repaired_words == 1


@pytest.mark.skipif(not _has_dict, reason="tr_TR sözlüğü yok")
def test_extract_formatting_differences_alone_do_not_split_word(tmp_path):
    # Kontrol: yalnız biçimlendirme farkı (kalın/punto/renk) docx2python'da
    # kelimeyi BÖLMEZ — bu daha önceki (yanlış) kök-neden varsayımının
    # düzeltilmiş hâlidir; regresyon olarak burada sabitlenir.
    from docx.shared import Pt, RGBColor

    document = docx.Document()
    p = document.add_paragraph()
    r1 = p.add_run("Kapsamadığ")
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0xFF, 0, 0)
    p.add_run("ı")
    path = tmp_path / "bicim_farki.docx"
    document.save(str(path))

    text = extract_docx(path)
    assert text == "Kapsamadığı"


def _make_docx(tmp_path, paragraphs):
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    path = tmp_path / "ornek.docx"
    document.save(str(path))
    return path


def test_extract_joins_paragraphs_with_blank_line(tmp_path):
    path = _make_docx(tmp_path, ["Birinci paragraf.", "İkinci paragraf."])
    text = extract_docx(path)
    assert text == "Birinci paragraf.\n\nİkinci paragraf."


def test_extract_drops_empty_paragraphs(tmp_path):
    path = _make_docx(tmp_path, ["Dolu.", "   ", "", "Yine dolu."])
    text = extract_docx(path)
    assert text == "Dolu.\n\nYine dolu."


def test_extract_output_is_chunkable(tmp_path):
    # Çıkarılan metin, boş-satır sınırından parçalanabilmeli (chunk.py ile uyum).
    from dilanaliz.chunk import chunk_text

    path = _make_docx(tmp_path, ["Bir.", "İki.", "Üç."])
    text = extract_docx(path)
    chunks = chunk_text(text, max_chars=4)
    assert [c.text for c in chunks] == ["Bir.", "İki.", "Üç."]


# --- Atlanan içeriklerin artık yakalandığını doğrula ----------------------

def test_extract_includes_table_cells(tmp_path):
    # python-docx'in atladığı tablo hücreleri docx2python ile okunmalı.
    document = docx.Document()
    document.add_paragraph("Gövde paragrafı.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Hücre A1"
    table.cell(0, 1).text = "Hücre B1"
    table.cell(1, 0).text = "Hücre A2"
    table.cell(1, 1).text = ""  # boş hücre atılmalı
    path = tmp_path / "tablo.docx"
    document.save(str(path))

    text, report = extract_docx_with_report(path)
    assert "Gövde paragrafı." in text
    for cell in ("Hücre A1", "Hücre B1", "Hücre A2"):
        assert cell in text
    assert report.tables == 1
    # Her dolu hücre + gövde paragrafı ayrı blok; boş hücre yok.
    assert text.count("\n\n") == 3


def test_extract_includes_header_and_footer(tmp_path):
    document = docx.Document()
    document.add_paragraph("Gövde.")
    section = document.sections[0]
    section.header.paragraphs[0].text = "ÜST BİLGİ metni"
    section.footer.paragraphs[0].text = "ALT BİLGİ metni"
    path = tmp_path / "hf.docx"
    document.save(str(path))

    text, report = extract_docx_with_report(path)
    assert "ÜST BİLGİ metni" in text
    assert "ALT BİLGİ metni" in text
    assert report.has_header_footer is True


def _png_bytes(width=2, height=2):
    """Geçerli, küçük bir RGB PNG üretir (python-docx boyut okuyabilsin)."""
    import struct
    import zlib

    def chunk(tag, data):
        body = tag + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)  # 8-bit RGB
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def test_report_counts_images_and_warns(tmp_path):
    # Görsel göm → sayılmalı ve uyarı üretmeli (içindeki yazı okunamaz).
    import io

    png = _png_bytes()
    document = docx.Document()
    document.add_paragraph("Görselli belge.")
    document.add_picture(io.BytesIO(png))
    path = tmp_path / "gorsel.docx"
    document.save(str(path))

    text, report = extract_docx_with_report(path)
    assert report.images >= 1
    assert any("görsel" in w.lower() for w in report.warnings)
    # Görsel yer tutucu işaretçisi analiz metnine sızmamalı.
    assert "media/" not in text
    assert "Görselli belge." in text
